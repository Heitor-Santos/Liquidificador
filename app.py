# -*- coding: utf-8 -*-
from itertools import permutations
from itertools import product
from itertools import islice
from random import shuffle
from docx import Document
from fpdf import FPDF
from pyfiglet import Figlet
import zipfile
import re
import os
import shutil

f = Figlet(font='small')
print(f.renderText('Liquidificador'))
print("Primeiro precisamos do nome do documento com o '.docx' no final")
print("Exemplos:    prova_3_ano.docx    provaFINAL.2019.docx    teste5serie.docx")
print("OBS: Se o documento não estiver na mesma pasta que o programa, digite o endereço inteiro")
print("Exemplos Windows:    C:\cliente\\fulano\\prova_3_ano.docx    D:\\pasta\\teste5serie.docx")
print("Exemplos Linux:  /home/fulano/Documentos/provaFINAL.2019.docx    /home/prvs/prova1.2019.docx")
print(" ")
nameDoc = input("Por favor digite aqui o nome do documento: ")
print("---------------------------------------------------------------------")
print(" ")
z = zipfile.ZipFile(nameDoc)
all_files = z.namelist()
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
#basicamente se tiver qualquer img no texto, ele salva em word/media/liquidImgs
liquidImgs = filter(lambda x: x.startswith('word/media/'), all_files)
j=0
for i in liquidImgs:
    image1 = z.open(i).read()
    z.extract(i, r'liquidImgs')
    j=j+1

doc = Document(nameDoc)
j=0
countPara=0
countImg=0
countAlt=0
countQuest=0
currDesc=[]
quests=[]
currAlts=[]
mapa={}
paras = doc.paragraphs
endHeader= False
print("Agora precisamos do gabarito da prova no formato NúmeroLetra-NúmeroLetra")
print("Exemplos: 1A-2B-3D-4A-5E     1B-2C-3B-4A-5C-6A-7B    1C-2C-3B")
print("OBS: As letras devem estar em maiúsculo")
print(" ")
feedback=input("Por favor digite aqui o gabarito: ")
print("------------------------------------------------------------------ ")
print(" ")
numExams=int(input("Por favor digite aqui o número de tipos de prova diferentes desejado: "))
print(" --------------------------------------------------------------")
print(" ")
answers={}
#separa o gabarito num array, uma posição para cada questão
feedback=feedback.split('-')
treatFeedback=[]
for i in feedback:
    #aqui deixa só a alternativa, sem o número da questão
    treatFeedback.append(i[len(i)-1]) 
for para in paras:
    #checa se tem uma imagem no parágrafo e a adiciona no enunciado da questão atuaç
    if 'graphicData' in para._p.xml:
        currDesc.append((1,countImg))
        countImg= countImg + 1
    
    elif re.match(r'^[A-za-z][).-]', para.text):
        currAlts.append((countPara,countAlt))
        countAlt= countAlt+1
        if countPara==len(paras)-1 or not re.match(r'^[A-za-z][).-]', paras[countPara+1].text):
           quests.append(currAlts)
           mapa[tuple(currAlts)]=currDesc
           answers[tuple(currAlts)]=treatFeedback[countQuest]
           currDesc=[]
           currAlts=[]
           countQuest+=1
    else: 
        if re.match(r'^(QUESTÃO|questão)( 01|1)',para.text)or re.match(r'^(01|1)[).-]', para.text):
            if(endHeader==False):
                endHeader=True
                header=currDesc
                currDesc=[]
                if re.match(r'^(QUESTÃO|questão)( 01)',para.text):
                    typeDescQuest=0
                elif re.match(r'^(QUESTÃO|questão)( 1)',para.text):
                    typeDescQuest=1
                elif re.match(r'^(01)', para.text):
                    if re.match(r'(01)[)]',para.text): typeDescQuest=2
                    elif re.match(r'(01)[.]',para.text): typeDescQuest=3
                    elif re.match(r'(01)[-]',para.text): typeDescQuest=4
                else:
                    if re.match(r'(1)[)]',para.text): typeDescQuest=5
                    elif re.match(r'(1)[.]',para.text): typeDescQuest=6
                    elif re.match(r'(1)[-]',para.text): typeDescQuest=7      
        currDesc.append((0,countPara))
    countPara = countPara + 1
mixOptions=[]
mixQuest=[]
permQuest=[]
finalQuest=[]
for i in quests: mixOptions.append(permutations(i))
mixQuest = islice(product(*mixOptions),1,10*numExams,10)
listMixQuest=[]
itemMixQuest=[]
for i in mixQuest:
    for j in i:
        itemMixQuest.append(list(j))
    listMixQuest.append(itemMixQuest)
    itemMixQuest=[] 
for i in listMixQuest:
    for j in i:
        shuffle(j)
for i in listMixQuest: shuffle(i)
countExam=0
pdf.cell(200, 10, txt="GABARITO "+nameDoc, ln=1, align="C")
for typeExam in listMixQuest:
    newExam = Document(nameDoc)
    newParas = newExam.paragraphs
    countPara=0
    countQuestion=1
    newFeedback=''
    for para in newParas:
        para.text=''
    for i in header:
        if i[0]==1:
            run = newParas[countPara].add_run()
            run.add_picture('liquidImgs/'+str(liquidImgs[i[1]]))
        else:
            newParas[countPara].text = paras[i[1]].text
        countPara+=1
    newParas[0].text= "TIPO "+str(countExam)+paras[0].text
    for question in typeExam:
        currDesc = mapa[tuple(sorted(question))]
        indexAns = ord(answers[tuple(sorted(question))])-65
        ans=sorted(question)[indexAns]
        newIndexAns = question.index(ans)
        newFeedback+=str(countQuestion)+"-"+chr(newIndexAns+65)+" "
        for i in currDesc:
            if i[0]==1:
                run = newParas[countPara].add_run()
                run.add_picture('liquidImgs/'+str(liquidImgs[i[1]]))
            else:
                newDesc = paras[i[1]].text
                if re.match(r'^(QUESTÃO|questão)', newDesc):
                    if typeDescQuest==0:
                        newDesc= newDesc[0:8]+ " "+str(countQuestion).zfill(2)+newDesc[11:]
                    else:
                        if re.match(r'[0-9]', newDesc[10]):
                            newDesc= newDesc[0:8]+ " "+str(countQuestion).zfill(2)+newDesc[11:]
                        else:
                            newDesc= newDesc[0:8]+ " "+str(countQuestion)+newDesc[10:]
                elif re.match(r'^[0-9][0-9][).-]', paras[i[1]].text):
                    if re.match(r'^[0-9][0-9][)]', paras[i[1]].text) and typeDescQuest==2:
                        newDesc= str(countQuestion).zfill(2)+newDesc[2:]
                    elif re.match(r'^[0-9][0-9][.]', paras[i[1]].text) and typeDescQuest==3:
                        newDesc= str(countQuestion).zfill(2)+newDesc[2:]
                    elif re.match(r'^[0-9][0-9][-]', paras[i[1]].text) and typeDescQuest==4:
                        newDesc= str(countQuestion).zfill(2)+newDesc[2:]
                elif re.match(r'^[0-9][).-]', paras[i[1]].text):
                    if re.match(r'^[0-9][)]', paras[i[1]].text) and typeDescQuest==5:
                        newDesc= str(countQuestion).zfill(2)+newDesc[1:]
                    elif re.match(r'^[0-9][.]', paras[i[1]].text) and typeDescQuest==6:
                        newDesc= str(countQuestion).zfill(2)+newDesc[1:]
                    elif re.match(r'^[0-9][-]', paras[i[1]].text) and typeDescQuest==7:
                        newDesc= str(countQuestion).zfill(2)+newDesc[1:]
                newDesc = str(newDesc)
                newParas[countPara].text = newDesc
            countPara+=1
        indexAlt=0
        for alt in question:
            newAlt= str(paras[alt[0]].text)
            newAlt = chr(indexAlt+65) + newAlt[1:]
            newParas[countPara].text = newAlt
            countPara+=1
            indexAlt+=1
        countQuestion+=1
    pdf.cell(200, 10, txt="TIPO "+str(countExam),ln=1, align="L")
    pdf.cell(200, 10, txt=newFeedback,ln=1, align="L")
    newExam.save(nameDoc[0:len(nameDoc)-5]+' TIPO '+str(countExam)+'.docx')
    countExam+=1
pdf.output("gabarito "+nameDoc+".pdf")
print("Provas e gabaritos gerados!")
if os.path.isdir('liquidImgs'):
    shutil.rmtree('liquidImgs')